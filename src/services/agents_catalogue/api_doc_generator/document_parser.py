"""
Document Parser for API Documentation Generator

This module handles parsing of various document formats including PDF, text,
and other document formats to extract content for API documentation generation.
"""

import os
import re
from pathlib import Path
from typing import Dict, Any, Optional, List

# Optional imports for different document formats
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.providers.logger import Logger


class DocumentParser:
    """Enhanced document parsing utilities for various formats"""
    
    def __init__(self, logger: Optional[Logger] = None):
        """Initialize document parser with optional logger"""
        self.logger = logger or Logger()
    
    def parse_document(self, file_path: str) -> str:
        """
        Parse document based on file extension and return extracted text
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content from the document
            
        Raises:
            Exception: If parsing fails or format is unsupported
        """
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        self.logger.info(f"Parsing document: {file_path} (type: {file_ext})")
        
        if file_ext == '.pdf':
            return self._parse_pdf_file(str(file_path))
        elif file_ext == '.txt':
            return self._parse_text_file(str(file_path))
        elif file_ext in ['.doc', '.docx']:
            return self._parse_docx_file(str(file_path))
        else:
            raise Exception(f"Unsupported document format: {file_ext}")
    
    def _parse_text_file(self, file_path: str) -> str:
        """Parse plain text files with encoding detection"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            if not content.strip():
                raise Exception(f"Text file is empty: {file_path}")
                
            self.logger.info(f"Successfully parsed text file: {len(content)} characters")
            return content
            
        except UnicodeDecodeError:
            # Fallback to latin-1 encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                    
                self.logger.info(f"Successfully parsed text file with latin-1 encoding: {len(content)} characters")
                return content
                
            except Exception as e:
                raise Exception(f"Error parsing text file {file_path}: {str(e)}")
        
        except Exception as e:
            raise Exception(f"Error parsing text file {file_path}: {str(e)}")
    
    def _parse_pdf_file(self, file_path: str) -> str:
        """Parse PDF files using pdfplumber with enhanced extraction"""
        if not PDF_AVAILABLE:
            raise Exception("PDF parsing not available. Please install pdfplumber: pip install pdfplumber")
        
        try:
            text_content = ""
            metadata = {}
            
            with pdfplumber.open(file_path) as pdf:
                # Extract metadata
                if pdf.metadata:
                    metadata = {
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                        'subject': pdf.metadata.get('Subject', ''),
                        'creator': pdf.metadata.get('Creator', ''),
                        'pages': len(pdf.pages)
                    }
                
                # Extract text from each page
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # Add page separator for better structure
                        text_content += f"\n=== PAGE {i+1} ===\n"
                        text_content += page_text.strip()
                        text_content += "\n"
                        
                        # Try to extract tables if they exist
                        tables = page.extract_tables()
                        if tables:
                            text_content += f"\n--- TABLES ON PAGE {i+1} ---\n"
                            for j, table in enumerate(tables):
                                text_content += f"Table {j+1}:\n"
                                for row in table:
                                    if row:
                                        text_content += " | ".join([cell or "" for cell in row]) + "\n"
                                text_content += "\n"
            
            if not text_content.strip():
                raise Exception(f"No text extracted from PDF: {file_path}")
            
            # Clean up the extracted text
            text_content = self._clean_extracted_text(text_content)
            
            self.logger.info(f"Successfully parsed PDF: {len(text_content)} characters, {metadata.get('pages', 0)} pages")
            
            # Add metadata as header if available
            if metadata:
                header = "=== DOCUMENT METADATA ===\n"
                for key, value in metadata.items():
                    if value:
                        header += f"{key.title()}: {value}\n"
                header += "\n=== DOCUMENT CONTENT ===\n"
                text_content = header + text_content
            
            return text_content
            
        except Exception as e:
            raise Exception(f"Error parsing PDF file {file_path}: {str(e)}")
    
    def _parse_docx_file(self, file_path: str) -> str:
        """Parse DOCX files using python-docx"""
        if not DOCX_AVAILABLE:
            raise Exception("DOCX parsing not available. Please install python-docx: pip install python-docx")
        
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract tables
            if doc.tables:
                text_content += "\n=== DOCUMENT TABLES ===\n"
                for i, table in enumerate(doc.tables):
                    text_content += f"\nTable {i+1}:\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        text_content += " | ".join(row_text) + "\n"
            
            if not text_content.strip():
                raise Exception(f"No text extracted from DOCX: {file_path}")
            
            # Clean up the extracted text
            text_content = self._clean_extracted_text(text_content)
            
            self.logger.info(f"Successfully parsed DOCX: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            raise Exception(f"Error parsing DOCX file {file_path}: {str(e)}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        # Remove trailing spaces
        lines = text.split('\n')
        cleaned_lines = [line.rstrip() for line in lines]
        text = '\n'.join(cleaned_lines)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def extract_document_structure(self, content: str) -> Dict[str, Any]:
        """
        Analyze document structure to identify sections, headers, and key information
        
        Args:
            content: Extracted document content
            
        Returns:
            Dictionary containing structured information about the document
        """
        structure = {
            'headers': [],
            'sections': {},
            'api_endpoints': [],
            'technical_terms': [],
            'urls': [],
            'statistics': {}
        }
        
        lines = content.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Detect headers (common patterns)
            if self._is_header_line(line):
                structure['headers'].append({
                    'text': line,
                    'line_number': i,
                    'level': self._get_header_level(line)
                })
                current_section = line
                structure['sections'][current_section] = []
            
            # Add content to current section
            if current_section and line not in [h['text'] for h in structure['headers']]:
                structure['sections'][current_section].append(line)
            
            # Extract URLs
            urls = re.findall(r'https?://[^\s<>"]+', line)
            structure['urls'].extend(urls)
            
            # Extract API endpoints (common patterns)
            api_patterns = [
                r'/api/[^\s<>"]+',
                r'/v\d+/[^\s<>"]+',
                r'POST\s+[^\s]+',
                r'GET\s+[^\s]+',
                r'PUT\s+[^\s]+',
                r'DELETE\s+[^\s]+'
            ]
            
            for pattern in api_patterns:
                matches = re.findall(pattern, line, re.IGNORECASE)
                structure['api_endpoints'].extend(matches)
        
        # Calculate statistics
        structure['statistics'] = {
            'total_lines': len(lines),
            'non_empty_lines': len([l for l in lines if l.strip()]),
            'total_characters': len(content),
            'sections_found': len(structure['sections']),
            'urls_found': len(structure['urls']),
            'api_endpoints_found': len(structure['api_endpoints'])
        }
        
        return structure
    
    def _is_header_line(self, line: str) -> bool:
        """Determine if a line is likely a header"""
        # Common header patterns
        header_patterns = [
            r'^[A-Z][A-Z\s]{3,}$',  # ALL CAPS
            r'^\d+\.\s+[A-Z]',       # Numbered headers
            r'^[A-Z][^.!?]*:$',      # Title with colon
            r'^={3,}.*={3,}$',       # Surrounded by equals
            r'^-{3,}.*-{3,}$',       # Surrounded by dashes
        ]
        
        return any(re.match(pattern, line) for pattern in header_patterns)
    
    def _get_header_level(self, line: str) -> int:
        """Determine header level (1-6)"""
        if re.match(r'^={3,}.*={3,}$', line):
            return 1
        elif re.match(r'^-{3,}.*-{3,}$', line):
            return 2
        elif re.match(r'^\d+\.\s+', line):
            return 3
        elif re.match(r'^[A-Z][A-Z\s]{3,}$', line):
            return 2
        elif re.match(r'^[A-Z][^.!?]*:$', line):
            return 4
        else:
            return 5 